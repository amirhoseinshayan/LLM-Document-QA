import tempfile
from io import BytesIO

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from docx import Document as DocxDocument
from rest_framework import status
from rest_framework.test import APIClient

from documents.models import Document, QuestionAnswer
from documents.services.document_processor import process_document
from documents.services.search_service import search_relevant_chunks


def create_test_docx_bytes() -> bytes:
    """
    Create DOCX file content in memory for tests.
    """
    file_stream = BytesIO()

    docx = DocxDocument()
    docx.add_heading("Test Literature Document", level=1)
    docx.add_paragraph(
        "This document is about literature, reading methods, Django, and document analysis."
    )
    docx.add_paragraph(
        "The system should extract this text, create chunks, and use them for search and question answering."
    )
    docx.save(file_stream)

    file_stream.seek(0)

    return file_stream.getvalue()


def create_test_uploaded_docx(filename: str = "test_document.docx") -> SimpleUploadedFile:
    """
    Create a Django uploaded DOCX file for API and model tests.
    """
    return SimpleUploadedFile(
        name=filename,
        content=create_test_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


class BaseDocumentTestCase(TestCase):
    """
    Base test case with isolated media storage.
    """

    def setUp(self):
        self.temp_media = tempfile.TemporaryDirectory()
        self.override_media = override_settings(MEDIA_ROOT=self.temp_media.name)
        self.override_media.enable()

    def tearDown(self):
        self.override_media.disable()
        self.temp_media.cleanup()

    def create_document_instance(self, title: str = "Test Document") -> Document:
        """
        Create and process a test document.
        """
        uploaded_file = create_test_uploaded_docx()

        document = Document.objects.create(
            title=title,
            file=uploaded_file,
        )

        process_document(document)
        document.refresh_from_db()

        return document


class DocumentProcessingTests(BaseDocumentTestCase):
    """
    Test DOCX processing and chunk creation.
    """

    def test_document_processing_extracts_text_and_creates_chunks(self):
        document = self.create_document_instance()

        self.assertTrue(document.is_processed)
        self.assertEqual(document.processing_error, "")
        self.assertIn("literature", document.full_text.lower())
        self.assertGreater(document.chunks.count(), 0)

    def test_reprocessing_recreates_chunks(self):
        document = self.create_document_instance()
        first_chunk_count = document.chunks.count()

        success = process_document(document)
        document.refresh_from_db()

        self.assertTrue(success)
        self.assertTrue(document.is_processed)
        self.assertGreater(document.chunks.count(), 0)
        self.assertEqual(document.chunks.count(), first_chunk_count)


class SearchServiceTests(BaseDocumentTestCase):
    """
    Test chunk search service.
    """

    def test_search_relevant_chunks_returns_results(self):
        document = self.create_document_instance()

        results = search_relevant_chunks(
            query="literature",
            top_k=5,
        )

        self.assertGreater(len(results), 0)
        self.assertEqual(results[0].chunk.document.id, document.id)

    def test_search_relevant_chunks_can_filter_by_document(self):
        document = self.create_document_instance()

        results = search_relevant_chunks(
            query="Django",
            top_k=5,
            document_id=document.id,
        )

        self.assertGreater(len(results), 0)

        for result in results:
            self.assertEqual(result.chunk.document.id, document.id)


class DocumentAPITests(BaseDocumentTestCase):
    """
    Test REST API endpoints.
    """

    def setUp(self):
        super().setUp()
        self.client = APIClient()

    def test_document_upload_api_processes_docx(self):
        uploaded_file = create_test_uploaded_docx("api_uploaded_document.docx")

        response = self.client.post(
            "/api/documents/",
            {
                "title": "API Uploaded Document",
                "file": uploaded_file,
            },
            format="multipart",
        )

        self.assertEqual(response.status_code, status.HTTP_201_CREATED, response.data)
        self.assertEqual(Document.objects.count(), 1)

        document = Document.objects.first()

        self.assertTrue(document.is_processed)
        self.assertEqual(document.processing_error, "")
        self.assertGreater(document.chunks.count(), 0)

    def test_document_list_api_returns_documents(self):
        self.create_document_instance()

        response = self.client.get("/api/documents/")

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertGreaterEqual(len(response.data), 1)

    def test_chunks_api_returns_chunks(self):
        self.create_document_instance()

        response = self.client.get("/api/chunks/")

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertGreaterEqual(len(response.data), 1)

    def test_search_api_returns_relevant_chunks(self):
        self.create_document_instance()

        response = self.client.post(
            "/api/search/",
            {
                "query": "literature",
                "top_k": 5,
            },
            format="json",
        )

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["query"], "literature")
        self.assertGreater(response.data["count"], 0)
        self.assertGreater(len(response.data["results"]), 0)

    @override_settings(LLM_PROVIDER="mock")
    def test_ask_api_generates_mock_answer_and_saves_history(self):
        self.create_document_instance()

        response = self.client.post(
            "/api/ask/",
            {
                "question": "What is this document about?",
                "top_k": 5,
            },
            format="json",
        )

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn("answer", response.data)
        self.assertIn("history_id", response.data)
        self.assertEqual(QuestionAnswer.objects.count(), 1)

    def test_history_api_returns_saved_records(self):
        document = self.create_document_instance()

        history = QuestionAnswer.objects.create(
            question="What is this document about?",
            answer="This is a test answer.",
        )
        history.related_documents.add(document)

        response = self.client.get("/api/history/")

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertGreaterEqual(len(response.data), 1)


class UserInterfaceTests(BaseDocumentTestCase):
    """
    Test Django Template user interface pages.
    """

    def test_home_page_loads(self):
        response = self.client.get("/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Document Question Answering System")

    def test_document_list_page_loads(self):
        response = self.client.get("/documents/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Documents")

    def test_document_upload_page_loads(self):
        response = self.client.get("/documents/upload/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Upload")

    def test_document_detail_page_loads(self):
        document = self.create_document_instance()

        response = self.client.get(f"/documents/{document.id}/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, document.title)

    def test_document_update_page_loads(self):
        document = self.create_document_instance()

        response = self.client.get(f"/documents/{document.id}/edit/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Edit Document")

    def test_document_delete_page_loads(self):
        document = self.create_document_instance()

        response = self.client.get(f"/documents/{document.id}/delete/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Delete Document")

    def test_search_page_loads(self):
        response = self.client.get("/search/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Search")

    @override_settings(LLM_PROVIDER="mock")
    def test_ask_page_loads(self):
        response = self.client.get("/ask/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Ask")

    def test_history_page_loads(self):
        response = self.client.get("/history/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "History")


class DocumentationTests(TestCase):
    """
    Test API documentation endpoints.
    """

    def test_swagger_ui_loads(self):
        response = self.client.get("/api/docs/")

        self.assertEqual(response.status_code, 200)

    def test_redoc_loads(self):
        response = self.client.get("/api/redoc/")

        self.assertEqual(response.status_code, 200)

    def test_openapi_schema_loads(self):
        response = self.client.get("/api/schema/")

        self.assertEqual(response.status_code, 200)